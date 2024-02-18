from docx import Document
from docx.shared import Pt
import os
import win32com.client


def get_run_formatting(run):
    """
    Extract formatting information from a Run object.
    """
    formatting = {
        'font': run.font.name,
        'size': run.font.size,
        'color': run.font.color.rgb,
        'bold': run.font.bold,
        'italic': run.font.italic,
        'underline': run.font.underline,
    }
    return formatting


def apply_formatting_to_run(run, formatting):
    """
    Apply formatting to a Run object.
    """
    run.font.name = formatting['font']
    run.font.size = formatting['size']
    run.font.color.rgb = formatting['color']
    run.font.bold = formatting['bold']
    run.font.italic = formatting['italic']
    run.font.underline = formatting['underline']


async def convert_to_pdf(docx_path, pdf_path):
    # Open Word application
    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False  # Don't show the Word application window

    # Open the document
    doc = word.Documents.Open(docx_path)

    # Save as PDF
    doc.SaveAs(pdf_path, FileFormat=17)  # FileFormat=17 for PDF

    # Close the document and Word application
    doc.Close()
    word.Quit()

# # Single
# def replace_text_in_docx(template_path, new_text, output_path):
#     # Load the template document
#     doc = Document(template_path)

#     # Replace text in all paragraphs
#     for para in doc.paragraphs:
#         if 'name' in para.text:
#             # Find and replace text in runs while preserving formatting
#             for run in para.runs:
#                 if 'name' in run.text:
#                     formatting = get_run_formatting(run)
#                     run.text = run.text.replace('name', new_text)
#                     apply_formatting_to_run(run, formatting)

#     # Save the modified document
#     doc.save(output_path)

# Multiple
async def replace_text_in_docx(template_path, to_replace , replace_with, output_path):
    # Load the template document
    doc = Document(template_path)

    # Replace text in all paragraphs
    for para in doc.paragraphs:
        if to_replace in para.text:
            # Find and replace text in runs while preserving formatting
            for run in para.runs:
                if to_replace in run.text:
                    formatting = get_run_formatting(run)
                    run.text = run.text.replace('name', replace_with)
                    apply_formatting_to_run(run, formatting)

    # Save the modified document
    doc.save(output_path)